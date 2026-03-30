from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "2563EB")
    rPr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "26")
    rPr.append(sz)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return paragraph

def style_doc(doc):
    """Set up document styles."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(15)
    font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    pf = style.paragraph_format
    pf.space_after = Pt(10)
    pf.line_spacing = 1.15

    for level, size, color in [
        ("Heading 1", 27, (0x17, 0x14, 0x12)),
        ("Heading 2", 21, (0x17, 0x14, 0x12)),
        ("Heading 3", 19, (0x3D, 0x3D, 0x3D)),
    ]:
        s = doc.styles[level]
        s.font.name = "Arial"
        s.font.size = Pt(size)
        s.font.color.rgb = RGBColor(*color)
        s.font.bold = True
        s.paragraph_format.space_before = Pt(18 if level != "Heading 1" else 24)
        s.paragraph_format.space_after = Pt(6)

def p(doc, text="", bold=False, italic=False, size=None, color=None):
    """Add a paragraph."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return para

def p_with_source(doc, text, source_name, source_url):
    """Add a paragraph with an inline hyperlink source."""
    para = doc.add_paragraph()
    run = para.add_run(text + "\n")
    run.font.size = Pt(15)
    src_run = para.add_run("Source: ")
    src_run.font.size = Pt(13)
    src_run.font.color.rgb = RGBColor(0x6B, 0x6B, 0x6B)
    add_hyperlink(para, source_name, source_url)
    return para

def stat_line(doc, stat_text, source_name, source_url):
    """Add a stat with source as a hyperlink."""
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(stat_text + " — ")
    run.font.size = Pt(15)
    add_hyperlink(para, source_name, source_url)
    return para

def note(doc, text):
    """Add a note/caveat paragraph."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(13)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x6B, 0x6B, 0x6B)
    return para

def source_list_item(doc, source_name, source_url):
    """Add a numbered source with hyperlink."""
    para = doc.add_paragraph(style="List Number")
    add_hyperlink(para, source_name, source_url)
    return para

# --- BUILD DOCUMENT ---

doc = Document()
style_doc(doc)

# Set margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)

# --- TITLE ---
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = title.add_run("[APP_NAME] Target Audience Brief")
run.font.size = Pt(37)
run.font.bold = True
run.font.color.rgb = RGBColor(0x17, 0x14, 0x12)

meta = doc.add_paragraph()
for label, value in [
    ("Status: ", "Living document"),
    ("  |  Last updated: ", "2026-03-17"),
    ("  |  Primary segment: ", "Women 45+, US (Gen X / Early Boomer)"),
]:
    r = meta.add_run(label)
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x6B, 0x6B, 0x6B)
    r = meta.add_run(value)
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x3D, 0x3D, 0x3D)

p(doc)  # spacer

# --- THE OPPORTUNITY ---
doc.add_heading("The Opportunity", level=1)

para = doc.add_paragraph()
para.add_run("Current earn/reward apps are built for younger audiences — gamified mechanics, confusing point systems, hype-driven marketing. Women 45+ are massively underserved: ").font.size = Pt(13)
r = para.add_run("59% of adults 50+ say technology is not designed with their age group in mind")
r.font.size = Pt(13)
r.bold = True
para.add_run(" (").font.size = Pt(13)
add_hyperlink(para, "AARP Tech Trends, 2024", "https://www.aarp.org/pri/topics/technology/internet-media-devices/2024-technology-trends-older-adults/")
para.add_run(").").font.size = Pt(13)

para = doc.add_paragraph()
para.add_run("They control ").font.size = Pt(13)
r = para.add_run("$15.2 trillion in global consumer spending")
r.font.size = Pt(13)
r.bold = True
para.add_run(" as part of Gen X (").font.size = Pt(13)
add_hyperlink(para, "NIQ / World Data Lab, 2025", "https://nielseniq.com/global/en/insights/report/2025/the-x-factor/")
para.add_run(") and influence ").font.size = Pt(13)
r = para.add_run("70–80% of all consumer purchasing decisions")
r.font.size = Pt(13)
r.bold = True
para.add_run(" (").font.size = Pt(13)
add_hyperlink(para, "NIQ, 2025", "https://nielseniq.com/global/en/insights/report/2025/the-x-factor/")
para.add_run("; originally ").font.size = Pt(13)
add_hyperlink(para, "BCG / Harvard Business Review, 2009", "https://hbr.org/2009/09/the-female-economy")
para.add_run("). Few products are designed for them.").font.size = Pt(13)

para = doc.add_paragraph()
para.add_run("A product built on trust, simplicity, and practical value can win this audience — and once won, they stay and refer friends. ").font.size = Pt(13)
r = para.add_run("88% of consumers trust word-of-mouth recommendations above all other forms of advertising")
r.font.size = Pt(13)
para.add_run(" (").font.size = Pt(13)
add_hyperlink(para, "Nielsen, 2021", "https://www.nielsen.com/insights/2021/beyond-martech-building-trust-with-consumers-and-engaging-where-sentiment-is-high/")
para.add_run(").").font.size = Pt(13)


# --- CORE MINDSET ---
doc.add_heading("Core Mindset", level=1)

doc.add_heading("Financial Reality", level=2)
p(doc, "This audience is financially cautious, not desperate. They feel the squeeze and want to ease it.")

doc.add_heading("Key signals", level=3)

stat_line(doc,
    "65% of women rank personal finances as their top source of stress and anxiety",
    "Laurel Road / HarrisX, 8th Annual Women's Financial Survey (2025)",
    "https://www.prnewswire.com/news-releases/women-taking-the-financial-reins-annual-survey-reveals-determination-amid-challenges-302393919.html")

stat_line(doc,
    "62% of women aged 50–64 say their personal finances are falling short of expectations",
    "AARP, \"She's the Difference\" Financial Security Survey (Dec 2025)",
    "https://www.aarp.org/pri/topics/work-finances-retirement/financial-security-retirement/financial-stress-older-women.html")

stat_line(doc,
    "51% of women 50+ feel less financially secure than a year ago",
    "AARP, \"She's the Difference\" Financial Security Survey (Dec 2025)",
    "https://www.aarp.org/pri/topics/work-finances-retirement/financial-security-retirement/shes-the-difference-financial-security/")

stat_line(doc,
    "65% of Gen X women worry about running out of money in retirement",
    "Allianz Life, 2025 Women Money Power Study (Sep 2025)",
    "https://www.allianzlife.com/about/newsroom/2025-Press-Releases/Gen-X-Nearing-Retirement-With-Worries")

stat_line(doc,
    "37% of Americans cannot cover a $400 emergency expense with cash",
    "Federal Reserve Board, SHED (2024, published May 2025)",
    "https://www.federalreserve.gov/publications/2024-economic-well-being-of-us-households-in-2023-expenses.htm")

p(doc, 'They are not looking for a "side hustle." They want small, reliable ways to ease everyday financial pressure — covering bills, handling unexpected expenses, stretching household budgets, feeling productive with their time.')


# --- BEHAVIORAL PROFILE ---
doc.add_heading("Behavioral Profile", level=2)

doc.add_heading("Daily Lifestyle Patterns", level=3)
p(doc, "Usage tends to occur during calm, routine moments in short bursts.")

para = doc.add_paragraph()
for window, time, context in [
    ("Morning", "7–9 AM", "coffee routine"),
    ("Early afternoon", "12–2 PM", "lunch break"),
    ("Evening", "8–10 PM", "wind-down"),
]:
    r = para.add_run(f"{window}: ")
    r.bold = True
    r.font.size = Pt(13)
    para.add_run(f"{time} ({context})\n").font.size = Pt(13)

p(doc, "Task tolerance: 5–10 minute activities. Simple micro-tasks. Clear value exchange.")
note(doc, "Note: These usage patterns are based on general behavioral research and internal product assumptions. Validate through in-app analytics post-launch.")


# --- TECHNOLOGY & APP BEHAVIOR ---
doc.add_heading("Technology & App Behavior", level=1)

doc.add_heading("Smartphone Ownership", level=2)

stat_line(doc,
    "91% of adults 50+ own a smartphone",
    "AARP, 2024 Technology Trends for Adults 50+",
    "https://www.aarp.org/pri/topics/technology/internet-media-devices/2024-technology-trends-older-adults/")

stat_line(doc,
    "90% of adults aged 50–64 own a smartphone; 78% of adults 65+ do",
    "Pew Research Center, Americans' Use of Mobile Technology (Jan 2024)",
    "https://www.pewresearch.org/internet/2024/01/31/americans-use-of-mobile-technology-and-home-broadband/")

p(doc, "Women 45+ are active smartphone users who favor practical, utility-focused apps: banking, shopping, health, streaming, and casual puzzle games.")


# --- SOCIAL MEDIA ---
doc.add_heading("Social Media Platforms", level=2)

# Table
table = doc.add_table(rows=6, cols=3)
table.style = "Light Grid Accent 1"

headers = ["Platform", "Adoption (50–64 / 65+)", "Key Role"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(13)

data = [
    ("Facebook", "74% / 57%", "Primary discovery & community"),
    ("YouTube", "85% / 64%", "Tutorials, reviews, entertainment"),
    ("Instagram", "40% / 19%", "Secondary inspiration"),
    ("TikTok", "30% / ~15%", "Growing but less dominant"),
    ("Pinterest", "~35% / 22%", "Shopping & lifestyle ideas"),
]
for i, (platform, adoption, role) in enumerate(data):
    for j, val in enumerate([platform, adoption, role]):
        cell = table.rows[i+1].cells[j]
        cell.text = val
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(13)

para = doc.add_paragraph()
para.add_run("\nSource: ").font.size = Pt(13)
add_hyperlink(para, "Pew Research Center, Americans' Social Media Use (2025)", "https://www.pewresearch.org/internet/2025/11/20/americans-social-media-use-2025/")


# --- GAMING ---
doc.add_heading("Entertainment & Game Preferences", level=1)

para = doc.add_paragraph()
para.add_run("This audience enjoys casual cognitive games, not fast-paced competitive gaming. Among gamers 50+, ").font.size = Pt(13)
r = para.add_run("women game daily at a higher rate than men (52% vs. 37%)")
r.bold = True
r.font.size = Pt(13)
para.add_run(" and 84% play on smartphones.").font.size = Pt(13)

doc.add_heading("Popular genres among gamers 50+", level=3)

genres = [
    ("Puzzle/logic games", "73%"),
    ("Card/tile games", "69%"),
    ("Word games", "58%"),
    ("Brain training", "37%"),
    ("Trivia/board games", "32%"),
]
for genre, pct in genres:
    para = doc.add_paragraph(style="List Bullet")
    r = para.add_run(f"{genre} — {pct}")
    r.font.size = Pt(13)

para = doc.add_paragraph()
para.add_run("Source: ").font.size = Pt(13)
add_hyperlink(para, 'AARP, "2023 Gamers 50+" (Apr 2023)', "https://www.aarp.org/pri/topics/social-leisure/activities-interests/2023-gamers-50-plus/")

p(doc, "Popular specific titles include Candy Crush and Solitaire (directly named in research). Words With Friends and Sudoku are consistent with genre preferences but not individually cited.")

para = doc.add_paragraph()
para.add_run("Source: ").font.size = Pt(13)
add_hyperlink(para, "AARP, Gaming Trends Among Older Americans (2020)", "https://www.aarp.org/pri/topics/technology/internet-media-devices/2020-gaming-trends-older-americans/")

para = doc.add_paragraph()
para.add_run("Supporting: Among Boomers, 73% prefer puzzle games and 55% prefer skill/chance games. 52% of Boomer women play video games vs. 46% of Boomer men. — ").font.size = Pt(13)
add_hyperlink(para, "ESA, Essential Facts (2025)", "https://www.theesa.com/annual-esa-study-reveals-video-games-universal-appeal-across-generations/")


# --- EARN APP ENGAGEMENT ---
doc.add_heading("Earn App Engagement", level=1)

doc.add_heading("Tasks That Work Well", level=2)
p(doc, "High-performing earning activities include receipt scanning for cashback, surveys with clear time/value exchange, shopping cashback, simple watch-to-earn content, and daily check-ins or streak bonuses.")
note(doc, "Note: These are based on industry patterns from earn/reward platforms (Swagbucks, Ibotta, InboxDollars) and internal product analysis. No credible public research breaks down earn-activity performance by age or gender for this specific demographic.")

para = doc.add_paragraph()
para.add_run("Gamification works best for this audience when it resembles casual routines and progress tracking, not competitive gameplay. Research on gamification for older adults finds the most effective elements are goal-setting, progress bars, rewards, points, and feedback — not leaderboards or competition. — ").font.size = Pt(13)
add_hyperlink(para, "Kappen et al., The Gerontologist (2021)", "https://pmc.ncbi.nlm.nih.gov/articles/PMC8437506/")

para = doc.add_paragraph()
para.add_run("Supporting: Leaderboards reduce social engagement among women regardless of competitive orientation. — ").font.size = Pt(13)
add_hyperlink(para, "Altmeyer et al., J. Computing in Higher Education (2025)", "https://link.springer.com/article/10.1007/s12528-025-09438-4")

doc.add_heading("Key Motivational Drivers", level=2)

p(doc, "1. Financial Relief", bold=True)
para = doc.add_paragraph()
para.add_run("Small earnings feel meaningful when framed as helping with everyday costs. This demographic strongly prefers clear monetary value — dollar balances and gift cards over complex point systems. Boomers and Gen X favor cash-back redemption at higher rates than younger cohorts. — ").font.size = Pt(13)
add_hyperlink(para, "Bankrate, Credit Card Rewards Survey (Nov 2024)", "https://www.bankrate.com/credit-cards/news/credit-card-rewards-survey/")

p(doc, "Effective: Dollar balances, gift cards, instant cashback.\nIneffective: Complex point systems, hidden reward tiers, long redemption delays.")

p(doc, "2. Productive Use of Time", bold=True)
p(doc, "Many are entering new life stages (empty nest, career shifts, retirement planning). Earning through an app provides a sense of productivity and control.")

p(doc, "3. Tangible Rewards", bold=True)
p(doc, "Frame everything around what she earned, not what she saved. This is not a coupon app, a deals platform, or a shopping companion. We pay people for their time and attention.")


# --- TRUST & SECURITY ---
doc.add_heading("Trust & Security", level=1)

p(doc, "Trust is the single biggest adoption barrier.", bold=True)

stat_line(doc,
    '81% of adults 50+ believe scams and fraud have reached "crisis" level. 91% recognize fraud can happen to anyone',
    "AARP, Fraud Awareness Report (2024)",
    "https://www.aarp.org/money/scams-fraud/aarp-fraud-awareness-report-2024/")

stat_line(doc,
    "Fraud losses reported by adults 60+ increased fourfold from ~$600M in 2020 to $2.4B in 2024. Adults 60+ are 5x more likely to lose money to tech support scams",
    'FTC, "Protecting Older Consumers 2024–2025" (Dec 2025)',
    "https://www.ftc.gov/system/files/ftc_gov/pdf/P144400-OlderAdultsReportDec2025.pdf")

doc.add_heading("Signals that cause immediate drop-off", level=3)
p(doc, "Unrealistic earning claims, aggressive notifications, countdown timers, confusing reward structures, poor visual design, hidden terms.")

para = doc.add_paragraph()
para.add_run("Source: FTC documents these as manipulative design patterns causing consumer harm. — ").font.size = Pt(13)
add_hyperlink(para, 'FTC, "Bringing Dark Patterns to Light" (Sep 2022)', "https://www.ftc.gov/reports/bringing-dark-patterns-light")

doc.add_heading("Trust-building mechanisms", level=3)
p(doc, "Transparent earning mechanics, clear privacy explanations, professional design, real testimonials, accessible customer support.")
p(doc, "If the app feels suspicious, users will leave immediately and not return.", bold=True)


# --- DESIGN PRINCIPLES ---
doc.add_heading("Design Principles", level=1)

doc.add_heading("Simplicity First", level=2)
p(doc, "Low cognitive load is essential. This audience prefers simple, predictable, and useful interfaces over flashy or heavily gamified ones.")
p(doc, "Design guidelines: Clear navigation, large typography, visible reward amounts, short task flows, tutorials when needed.")

para = doc.add_paragraph()
para.add_run("Source: ").font.size = Pt(13)
add_hyperlink(para, "Nielsen Norman Group, Usability for Senior Citizens", "https://www.nngroup.com/articles/usability-for-senior-citizens/")
para.add_run(" — 87 design guidelines emphasizing predictable layouts, minimum 14pt font, and consistent interactions.").font.size = Pt(13)

para = doc.add_paragraph()
para.add_run("Supporting: Reducing interface complexity helps older adults find features as quickly as younger adults. — ").font.size = Pt(13)
add_hyperlink(para, "ACM CHI 2024", "https://dl.acm.org/doi/10.1145/3613904.3642796")


# --- HABIT & RETENTION ---
doc.add_heading("Habit & Retention Mechanics", level=1)

p(doc, "This audience responds strongly to routine-based engagement loops.")
p(doc, "Effective patterns: Daily check-in bonuses, streak rewards, weekly earning summaries, seasonal reward campaigns.")
note(doc, "Note: These are product design principles based on general retention research and industry patterns. No credible public source validates these specific mechanics for women 45+ in earn/reward apps. Validate through A/B testing post-launch.")

para = doc.add_paragraph()
para.add_run("Women score higher on collaborative, social, and care-taking gamification types; men score higher on achiever and competitive types. — ").font.size = Pt(13)
add_hyperlink(para, "Tondello et al., Intl Journal of Human-Computer Interaction (2024)", "https://www.tandfonline.com/doi/full/10.1080/10447318.2024.2446498")


# --- SOCIAL & COMMUNITY ---
doc.add_heading("Social & Community", level=1)

p(doc, "They prefer trusted circles over public competition.")
p(doc, "Effective: Referral rewards, share deals with friends, small trusted communities.\nIneffective: Leaderboards, competitive rankings, public bragging mechanics.")

para = doc.add_paragraph()
para.add_run("88% of consumers trust recommendations from people they know above all other forms of advertising. — ").font.size = Pt(13)
add_hyperlink(para, "Nielsen, Global Trust in Advertising (2021)", "https://www.nielsen.com/insights/2021/beyond-martech-building-trust-with-consumers-and-engaging-where-sentiment-is-high/")

para = doc.add_paragraph()
para.add_run("Among adults 50+, 17% cite recommendations from friends and family as the primary factor when deciding what to buy. — ").font.size = Pt(13)
add_hyperlink(para, "AARP, 2023 Technology Trends for Adults 50+", "https://www.aarp.org/pri/topics/technology/internet-media-devices/2023-technology-trends-older-adults/")


# --- BRAND POSITIONING ---
doc.add_heading("Brand Positioning", level=1)

p(doc, "The brand must feel trustworthy, competent, respectful, practical, calm, and professional.")
p(doc, 'Messaging should frame earning as smart and responsible behavior.\n\nEffective tone: "You\'re smart for doing this."\nAvoid: "Easy money," hype language, get-rich messaging.')


# --- STRATEGIC OPPORTUNITY ---
doc.add_heading("Strategic Opportunity", level=1)

p(doc, "This audience represents a large and underserved market.", bold=True)

stat_line(doc,
    "Gen X is projected to drive $15.2 trillion in global consumer spending in 2025",
    'NIQ / World Data Lab, "The X Factor" (Jul 2025)',
    "https://nielseniq.com/global/en/insights/report/2025/the-x-factor/")

stat_line(doc,
    "Women influence 70–80% of all consumer purchasing decisions",
    "NIQ (2025); originally BCG / Harvard Business Review (2009)",
    "https://hbr.org/2009/09/the-female-economy")

stat_line(doc,
    "Women 50+ control 95% of household purchasing decisions",
    "AARP (citing Nielsen and U.S. Consumer Expenditure Survey)",
    "https://advertise.aarp.org/50-insights/ignoring-50-plus-women")

stat_line(doc,
    "59% of adults 50+ say technology is not designed with their age group in mind",
    "AARP, Technology Trends (2022–2024)",
    "https://www.aarp.org/pri/topics/technology/internet-media-devices/2024-technology-trends-older-adults/")

p(doc, "A product designed around trust, simplicity, and practical financial value can achieve strong retention and word-of-mouth growth.")


# --- CORE PRODUCT PRINCIPLES ---
doc.add_heading("Core Product Principles", level=1)

p(doc, "The app should feel like a trustworthy budgeting ally that rewards everyday actions.", bold=True)

principles = [
    "Simple earning",
    "Clear rewards",
    "Strong trust signals",
    "Routine-friendly engagement",
    "Referral-driven growth",
]
for principle in principles:
    para = doc.add_paragraph(style="List Bullet")
    para.add_run(principle).font.size = Pt(13)


# --- COMPLETE SOURCE LIST ---
doc.add_heading("Complete Source List", level=1)

sources = [
    ("Laurel Road / HarrisX, 8th Annual Women's Financial Survey (2025)", "https://www.prnewswire.com/news-releases/women-taking-the-financial-reins-annual-survey-reveals-determination-amid-challenges-302393919.html"),
    ('AARP, "She\'s the Difference" — Financial Stress Among Women 50–64 (Dec 2025)', "https://www.aarp.org/pri/topics/work-finances-retirement/financial-security-retirement/financial-stress-older-women.html"),
    ('AARP, "She\'s the Difference" — Financial Anxiety Among Women 50+ (Dec 2025)', "https://www.aarp.org/pri/topics/work-finances-retirement/financial-security-retirement/shes-the-difference-financial-security/"),
    ("Allianz Life, 2025 Women Money Power Study (Sep 2025)", "https://www.allianzlife.com/about/newsroom/2025-Press-Releases/Gen-X-Nearing-Retirement-With-Worries"),
    ("Federal Reserve Board, SHED — Economic Well-Being of U.S. Households (2024)", "https://www.federalreserve.gov/publications/2024-economic-well-being-of-us-households-in-2023-expenses.htm"),
    ("Pew Research Center, Americans' Use of Mobile Technology (Jan 2024)", "https://www.pewresearch.org/internet/2024/01/31/americans-use-of-mobile-technology-and-home-broadband/"),
    ("Pew Research Center, Americans' Social Media Use (2025)", "https://www.pewresearch.org/internet/2025/11/20/americans-social-media-use-2025/"),
    ("AARP, 2024 Technology Trends for Adults 50+", "https://www.aarp.org/pri/topics/technology/internet-media-devices/2024-technology-trends-older-adults/"),
    ("Deloitte, Connected Consumer Survey (2023)", "https://www.deloitte.com/us/en/insights/industry/telecommunications/connectivity-mobile-trends-survey/2023/connectivity-mobile-trends-survey-full-report.html"),
    ('AARP, "2023 Gamers 50+" (Apr 2023)', "https://www.aarp.org/pri/topics/social-leisure/activities-interests/2023-gamers-50-plus/"),
    ("AARP, Gaming Trends Among Older Americans (2020)", "https://www.aarp.org/pri/topics/technology/internet-media-devices/2020-gaming-trends-older-americans/"),
    ("Entertainment Software Association, Essential Facts (2025)", "https://www.theesa.com/annual-esa-study-reveals-video-games-universal-appeal-across-generations/"),
    ("AARP, Fraud Awareness Report (2024)", "https://www.aarp.org/money/scams-fraud/aarp-fraud-awareness-report-2024/"),
    ('FTC, "Protecting Older Consumers 2024–2025" (Dec 2025)', "https://www.ftc.gov/system/files/ftc_gov/pdf/P144400-OlderAdultsReportDec2025.pdf"),
    ('FTC, "Bringing Dark Patterns to Light" (Sep 2022)', "https://www.ftc.gov/reports/bringing-dark-patterns-light"),
    ("Nielsen Norman Group, Usability for Senior Citizens", "https://www.nngroup.com/articles/usability-for-senior-citizens/"),
    ("Kappen et al., The Gerontologist (2021)", "https://pmc.ncbi.nlm.nih.gov/articles/PMC8437506/"),
    ("Altmeyer et al., J. Computing in Higher Education (2025)", "https://link.springer.com/article/10.1007/s12528-025-09438-4"),
    ("Tondello et al., Intl Journal of Human-Computer Interaction (2024)", "https://www.tandfonline.com/doi/full/10.1080/10447318.2024.2446498"),
    ("ACM CHI 2024", "https://dl.acm.org/doi/10.1145/3613904.3642796"),
    ('NIQ / World Data Lab, "The X Factor" (Jul 2025)', "https://nielseniq.com/global/en/insights/report/2025/the-x-factor/"),
    ('Harvard Business Review, "The Female Economy" (Sep 2009)', "https://hbr.org/2009/09/the-female-economy"),
    ("AARP, Ignoring 50+ Women", "https://advertise.aarp.org/50-insights/ignoring-50-plus-women"),
    ("Nielsen, Global Trust in Advertising (2021)", "https://www.nielsen.com/insights/2021/beyond-martech-building-trust-with-consumers-and-engaging-where-sentiment-is-high/"),
    ("Bankrate, Credit Card Rewards Survey (Nov 2024)", "https://www.bankrate.com/credit-cards/news/credit-card-rewards-survey/"),
    ("AARP, 2023 Technology Trends for Adults 50+", "https://www.aarp.org/pri/topics/technology/internet-media-devices/2023-technology-trends-older-adults/"),
]

for i, (name, url) in enumerate(sources):
    para = doc.add_paragraph()
    r = para.add_run(f"{i+1}. ")
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x6B, 0x6B, 0x6B)
    add_hyperlink(para, name, url)

# Save
output_path = "/Users/markus/Documents/earnapp/docs/Target_Audience_Brief.docx"
doc.save(output_path)
print(f"Saved to {output_path}")
